from email.mime import application
import html
import logging
import re

logger = logging.getLogger(__name__)


def remove_html_tags(text: str) -> str:
    """从文本中移除HTML标签"""
    if not text:
        return ""

    # 首先反转义HTML实体
    text = html.unescape(text)

    # 移除HTML标签
    clean = re.compile("<.*?>")
    text = re.sub(clean, "", text)

    return text


async def extract_text_content(file_path: str, mime_type: str) -> str:
    """使用健壮的处理程序从文件路径提取文本内容，
    这镜像了增强文档服务的提取逻辑，以便可以在知识库摄取和简历筛选之间重用
    """
    try:
        if mime_type in ("text/plain", "text/markdown"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif mime_type == "application/pdf":
            text = ""
            try:
                import PyPDF2
            except:
                logger.error(f"PyPDF2不可用于PDF提取: {e}")
                return ""
            with open(file_path, "rb") as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_texts = page.extract_text() or ""
                    if page_texts:
                        text += page_texts + "\n"
            return text.strip()
        elif mime_type in (
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ):
            if (
                mime_type
                == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                try:
                    import docx2txt

                    extracted = docx2txt.process(file_path) or ""
                    if extracted.strip():
                        logger.info("使用docx2txt提取DOCX内容")
                        return extracted.strip()
                    else:
                        logger.info("docx2txt返回空;回退到python-docx")

                except Exception as e:
                    logger.info(f"docx2txt不可用或失败({e})；回退到python-docx")
                try:
                    from docx import Document as DocxDocument

                    doc = DocxDocument(file_path)
                    parts = []
                    for p in doc.paragraphs:
                        if p.text and p.text.strip():
                            parts.append(p.text)

                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text and cell.text.strip():
                                    parts.append(cell.text)

                    return "\n".join(parts).strip()
                except Exception as e:
                    logger.error(f"python-docx提取内容失败{e}")
        else:

            logger.warning(f"不支持的文件类型: {mime_type}")
            return ""
    except Exception as e:
        logger.error(f"从{file_path}提取文件时出错")
